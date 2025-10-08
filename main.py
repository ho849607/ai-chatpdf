# legal_agent_streamlit_app.py
# -*- coding: utf-8 -*-
import io
import re
from dataclasses import dataclass, asdict
from typing import List, Dict, Any, Tuple

import streamlit as st

# 선택 라이브러리 (없어도 앱은 돌아감)
try:
    from docx import Document  # python-docx
except Exception:
    Document = None

try:
    import pandas as pd
except Exception:
    pd = None

# =====================[ UI 기본 설정 ]=====================
st.set_page_config(page_title="SpeciAI — 법률 의견서 에이전트", page_icon="⚖️", layout="wide")
st.title("⚖️ SpeciAI — 법률 의견서 에이전트 (Streamlit Demo v2)")

# 데모 대시보드 이미지 노출
DASHBOARD_IMG = "/mnt/data/4fa0f4cb-0166-4cbe-ad02-6cfdd2cd101f.png"
with st.expander("📷 데모 대시보드 이미지 보기/숨기기", expanded=False):
    try:
        st.image(DASHBOARD_IMG, caption="샘플 대시보드 UI (참고용)", use_container_width=True)
    except Exception as e:
        st.info("이미지를 불러올 수 없습니다. 경로 확인 필요.")
        st.exception(e)

st.markdown("---")

# =====================[ 데이터 모델 ]=====================
@dataclass
class BusinessInfo:
    사업체명: str = ""
    업종: str = ""
    국가: str = "대한민국"
    웹사이트: str = ""
    담당자명: str = ""
    연락처: str = ""
    특징: str = ""

@dataclass
class Engagement:
    목적: str = ""       # 예: 개인정보 처리방침 개정 검토
    법률분야: str = ""   # 개인정보/전자금융/노동/지식재산/공정거래/계약/기타
    추가설명: str = ""

@dataclass
class OpinionDoc:
    제목: str = ""
    본문: str = ""
    리스크라벨: List[Tuple[str, str, str]] = None  # (이슈, 심각도, 근거)

# =====================[ 규칙 기반 도우미 ]=====================
# (keyword, severity, label, rationale)
RISK_KEYWORDS = [
    ("민감정보", "심각", "개인정보보호법(민감정보)", "민감정보는 별도 동의 및 암호화·접근통제가 필요합니다."),
    ("주민등록번호", "심각", "개인정보보호법(고유식별정보)", "고유식별정보는 엄격한 보호·마스킹이 요구됩니다."),
    ("전자금융", "주의", "전자금융거래법", "전자지급/PG 등은 보안인증·사고책임 규제가 있습니다."),
    ("결제", "주의", "전자금융거래법/여신전문금융", "결제·수납은 결제대행·분할납부 등 규제 검토 필요."),
    ("쿠키", "주의", "통신비밀보호/개인정보", "행태정보 수집·광고 쿠키는 고지·동의·옵트아웃 요구."),
    ("미성년자", "심각", "청소년보호/개인정보", "14세 미만 동의·연령확인·유해매체 차단 필요."),
    ("노동자", "주의", "근로기준법/산안법", "근로시간·휴게·연장수당·산안 준수 필요."),
    ("하도급", "주의", "하도급법/공정거래", "우월적 지위 남용 금지, 서면발급 의무."),
    ("상표", "정보", "상표법", "표장 유사·식별력·선사용·선행조사 권장."),
    ("특허", "정보", "특허법", "신규성·진보성·명세서 기재요건 검토."),
    ("위치정보", "주의", "위치정보보호법", "위치사업 신고·동의·보관기간 제한."),
    ("클라우드", "정보", "전자문서/정보보호", "국외 이전·가명처리·접근통제 정책 검토."),
]

def risk_scan(text: str) -> List[Tuple[str, str, str]]:
    """아주 단순한 키워드 기반 리스크 라벨링(데모)."""
    results = []
    low = text.lower()
    for kw, sev, label, why in RISK_KEYWORDS:
        if kw.lower() in low:
            results.append((label, sev, why))
    uniq = []
    seen = set()
    for item in results:
        if item not in seen:
            uniq.append(item)
            seen.add(item)
    return uniq

def suggest_edits(text: str) -> List[str]:
    """간단한 문체/구조 제안 (LLM 연결 전 임시)."""
    edits = []
    if len(text.strip()) < 400:
        edits.append("문서 길이가 짧습니다. 배경-사실관계-쟁점-검토-결론의 5단 구성으로 확장하세요.")
    if "할 수 있다" in text:
        edits.append("표현이 모호합니다. '할 수 있다' 대신 허용요건·책임주체를 특정하세요.")
    if text.count("…") > 0:
        edits.append("생략부호(…) 대신 정확한 인용 또는 구체적 사실을 기재하세요.")
    if "개인정보" in text and "동의" not in text:
        edits.append("개인정보 처리의 동의/법적 근거(제15조 등)를 명시하세요.")
    if "전자금융" in text and "보안" not in text:
        edits.append("전자금융은 보안·인증·사고책임 분담 규정 언급이 필요합니다.")
    if "근로" in text and "근로시간" not in text:
        edits.append("노동 이슈는 근로시간·휴게·연장수당·서면계약 필수사항을 점검하세요.")
    return edits or ["큰 오류는 없으나, 판례·유권해석 인용으로 설득력을 높이세요."]

def multi_agent_outline(domain: str, purpose: str, extras: Dict[str, Any]) -> Dict[str, str]:
    """도메인별 자동 뼈대(멀티-에이전트 합성 데모)."""
    base = {
        "배경": f"{purpose} 관련 사실관계 요약 및 사업모델 설명.",
        "쟁점": "관련 법령/가이드라인 대비 위법/위험 포인트 정리.",
        "검토": "법령 조문·판례·감독지침·관행 순으로 검토.",
        "개선방안": "위험 완화 대안 및 실행 체크리스트.",
        "결론": "가능/불가/조건부 가능 등 최종 의견.",
    }
    if domain == "개인정보":
        base["검토"] = "개인정보보호법·시행령/고시·국외이전·가명정보 처리 가능성."
        base["개선방안"] = "최소수집·목적외 이용 금지·보관기간·암호화/접근통제·마스킹 정책."
    elif domain == "전자금융":
        base["검토"] = "전자금융거래법·여전법·PG/선불/지급수단·정산흐름."
        base["개선방안"] = "인증/보안 아키텍처·사고책임 분담·약관 고지·정산/보관."
    elif domain == "노동":
        base["검토"] = "근로기준법·근로계약 필수기재·연장/야간/휴일수당·52시간제."
        base["개선방안"] = "근로시간 시스템·휴게·임금명세서·4대보험·산안 리스크."
    elif domain == "지식재산":
        base["검토"] = "특허/상표/저작권 요건·권리범위·침해 가능성."
        base["개선방안"] = "선행조사·출원 전 공개 금지·표장/디자인 가이드."
    elif domain == "공정거래":
        base["검토"] = "하도급/대리점/플랫폼 공정화·우월적 지위 남용 금지."
        base["개선방안"] = "표준계약서·서면교부·단가조정·보복금지 모니터링."
    return base

# 인라인 하이라이트용 색상
HIGHLIGHT_COLORS = {"심각": "#fee2e2", "주의": "#fff7ed", "정보": "#e2e8f0"}

def highlight_risks_html(text: str) -> str:
    """키워드에 색 배경을 주는 간단한 HTML 마킹."""
    html = st.html_escape(text) if hasattr(st, "html_escape") else text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
    # 긴 키워드가 먼저 매칭되도록 길이순 정렬
    sorted_kw = sorted(RISK_KEYWORDS, key=lambda x: len(x[0]), reverse=True)
    for kw, sev, label, why in sorted_kw:
        color = HIGHLIGHT_COLORS.get(sev, "#e2e8f0")
        pattern = re.escape(kw)
        html = re.sub(
            pattern,
            f'<span style="background:{color}; padding:0 4px; border-radius:4px;" title="{label} — {why}">{kw}</span>',
            html,
            flags=re.IGNORECASE,
        )
    return f"<div style='line-height:1.7'>{html}</div>"

# 에이전트 역제안(도메인별 체크리스트)
DOMAIN_SUGGESTIONS = {
    "개인정보": [
        "수집항목·처리목적·보유기간 표로 정리",
        "민감정보/고유식별정보 분리 저장 및 접근통제",
        "국외이전 여부/경로/보관지역 명시",
        "파기절차·파기기록 로깅 설계",
    ],
    "전자금융": [
        "정산흐름(예치·분리보관) 도식화",
        "이중인증(FIDO/OTP) 적용 범위 결정",
        "사고 책임·면책 요건 약관 반영",
        "모의해킹/보안 점검 주기 설정",
    ],
    "노동": [
        "근로시간 시스템(주/연장/야간) 설정",
        "근로계약서 필수기재 항목 점검",
        "임금명세서 템플릿 적용",
        "산업안전 체크리스트 도입",
    ],
    "지식재산": [
        "선행조사 리포트 첨부",
        "출원전 공개 금지 정책 공지",
        "상표/디자인 가이드 배포",
        "오픈소스 라이선스 검토",
    ],
    "공정거래": [
        "표준계약서 채택(서면교부)",
        "단가조정·보복금지 조항 명문화",
        "대리점·플랫폼 수수료 투명화",
        "우월적 지위 남용 감시 체계",
    ],
    "계약": [
        "책임제한·손해배상·면책 조항 구체화",
        "준거법·관할·분쟁해결 절차 명시",
        "SLA/성능지표/위반시 구제수단",
        "비밀정보 범위와 예외 명확화",
    ],
    "기타": ["관련 업권 가이드라인 수집", "업계 표준/관행 대비표 작성"],
}

# =====================[ 세션 상태 ]=====================
if "biz" not in st.session_state:
    st.session_state.biz = BusinessInfo()
if "eng" not in st.session_state:
    st.session_state.eng = Engagement()
if "extras" not in st.session_state:
    st.session_state.extras = {}
if "opinion" not in st.session_state:
    st.session_state.opinion = OpinionDoc()
if "log" not in st.session_state:
    st.session_state.log = []  # 간단 작업 로그

def log(msg: str):
    st.session_state.log.append(msg)

# =====================[ 사이드바: 흐름 선택 ]=====================
st.sidebar.header("작업 흐름")
flow = st.sidebar.radio(
    "무엇을 하시겠어요?",
    ["① 기본정보 수집", "② 도메인 추가질문", "③ 초안 리뷰", "④ 의견서 생성", "⑤ 출력/발송"],
    index=0
)

with st.sidebar.expander("📝 작업 로그", expanded=False):
    if st.session_state.log:
        for i, m in enumerate(st.session_state.log[-20:], 1):
            st.write(f"{i}. {m}")
    else:
        st.caption("로그가 여기에 표시됩니다.")

# =====================[ ① 기본정보 수집 ]=====================
if flow == "① 기본정보 수집":
    st.subheader("① 사업체/고객사 기본정보")
    with st.form("basic_form"):
        col1, col2 = st.columns(2)
        with col1:
            name = st.text_input("사업체명", st.session_state.biz.사업체명)
            industry = st.text_input("업종", st.session_state.biz.업종)
            country = st.text_input("국가", st.session_state.biz.국가 or "대한민국")
            site = st.text_input("웹사이트(URL)", st.session_state.biz.웹사이트)
        with col2:
            pic = st.text_input("담당자명", st.session_state.biz.담당자명)
            contact = st.text_input("연락처/이메일", st.session_state.biz.연락처)
            desc = st.text_area("특징/요약(선택)", st.session_state.biz.특징, height=120)
        submitted = st.form_submit_button("저장")
    if submitted:
        st.session_state.biz = BusinessInfo(name, industry, country, site, pic, contact, desc)
        st.success("기본정보 저장 완료")
        log("기본정보 저장")
        st.json(asdict(st.session_state.biz))

    st.markdown("—")
    st.subheader("자문 개요")
    with st.form("engagement_form"):
        purpose = st.text_input("자문 목적 (예: 개인정보 처리방침 개정 검토)", st.session_state.eng.목적)
        domain = st.selectbox("질문하는 법률 분야", ["개인정보","전자금융","노동","지식재산","공정거래","계약","기타"], index=0)
        note = st.text_area("추가 설명", st.session_state.eng.추가설명, height=120)
        ok = st.form_submit_button("저장")
    if ok:
        st.session_state.eng = Engagement(purpose, domain, note)
        st.success("자문 개요 저장 완료")
        log(f"자문 개요 저장 — 분야: {domain}")
        st.json(asdict(st.session_state.eng))

# =====================[ ② 도메인 추가질문 + 역제안 ]=====================
elif flow == "② 도메인 추가질문":
    st.subheader("② 도메인별 필수 추가정보")
    domain = st.session_state.eng.법률분야 or st.selectbox("법률 분야 선택", ["개인정보","전자금융","노동","지식재산","공정거래","계약","기타"])
    extras: Dict[str, Any] = st.session_state.extras or {}

    with st.form("extras_form"):
        if domain == "개인정보":
            c1, c2 = st.columns(2)
            with c1:
                extras["개인정보_수집항목"] = st.text_area("수집 항목", extras.get("개인정보_수집항목",""))
                extras["개인정보_처리목적"] = st.text_area("처리 목적", extras.get("개인정보_처리목적",""))
            with c2:
                extras["보유기간"] = st.text_input("보유기간 (예: 1년 보관 후 지체없이 파기)", extras.get("보유기간",""))
                extras["국외이전"] = st.selectbox("국외이전 여부", ["아니오","예"], index=0 if extras.get("국외이전","아니오")=="아니오" else 1)
        elif domain == "전자금융":
            extras["결제유형"] = st.multiselect("결제유형", ["PG","선불충전","송금","BNPL","구독결제"], default=extras.get("결제유형",[]))
            extras["인증수준"] = st.selectbox("인증 수준", ["기본","강화"], index=0 if extras.get("인증수준","기본")=="기본" else 1)
            extras["정산주체"] = st.text_input("정산 주체/흐름", extras.get("정산주체",""))
        elif domain == "노동":
            extras["근로형태"] = st.selectbox("근로형태", ["정규직","계약직","프리랜서","인턴"],
                                         index={"정규직":0,"계약직":1,"프리랜서":2,"인턴":3}.get(extras.get("근로형태","정규직"),0))
            extras["근로시간"] = st.text_input("근로시간 (예: 주40시간)", extras.get("근로시간",""))
            extras["임금체계"] = st.text_input("임금체계/수당", extras.get("임금체계",""))
        elif domain == "지식재산":
            extras["이슈"] = st.multiselect("이슈", ["특허","상표","저작권","영업비밀"], default=extras.get("이슈",[]))
            extras["선행조사"] = st.selectbox("선행조사 여부", ["미실시","진행중","완료"],
                                          index={"미실시":0,"진행중":1,"완료":2}.get(extras.get("선행조사","미실시"),0))
        elif domain == "공정거래":
            extras["거래유형"] = st.multiselect("거래유형", ["하도급","대리점","플랫폼","유통"], default=extras.get("거래유형",[]))
            extras["우월적지위"] = st.selectbox("우월적 지위 의심", ["아니오","예"], index=0 if extras.get("우월적지위","아니오")=="아니오" else 1)
        else:
            extras["요청사항"] = st.text_area("기타 요청/배경", extras.get("요청사항",""))

        st.markdown("—")
        st.markdown("### 🤝 에이전트 역제안")
        sug = DOMAIN_SUGGESTIONS.get(domain, DOMAIN_SUGGESTIONS["기타"])
        checked = []
        cols = st.columns(2)
        for i, s in enumerate(sug):
            with cols[i % 2]:
                if st.checkbox(s, key=f"sugg_{i}", value=False):
                    checked.append(s)

        saved = st.form_submit_button("저장")
    if saved:
        extras["역제안"] = checked
        st.session_state.extras = extras
        st.success("도메인 추가정보/역제안 저장 완료")
        log(f"도메인 정보 저장 — {domain}, 역제안 {len(checked)}건")
        st.json(st.session_state.extras)

# =====================[ ③ 초안 리뷰 ]=====================
elif flow == "③ 초안 리뷰":
    st.subheader("③ 사용자가 작성한 의견서 초안 — 자동 리뷰")

    # 파일 업로드(.txt/.md)
    up = st.file_uploader("초안 파일 업로드 (.txt / .md)", type=["txt","md"])
    init_text = st.session_state.opinion.본문 or ""
    if up is not None:
        init_text = up.read().decode("utf-8", errors="ignore")
        log(f"초안 파일 업로드 — {up.name}")

    draft = st.text_area("초안 텍스트", height=280, value=init_text)

    col = st.columns([1,1,1,5])
    with col[0]:
        run = st.button("리뷰 실행")
    with col[1]:
        save_btn = st.button("현재 텍스트 저장")
    with col[2]:
        clear_btn = st.button("지우기")

    if save_btn:
        st.session_state.opinion.본문 = draft
        st.success("초안 텍스트 저장")
        log("초안 텍스트 저장")

    if clear_btn:
        st.session_state.opinion.본문 = ""
        st.warning("초안 텍스트를 비웠습니다.")
        log("초안 텍스트 초기화")

    if run:
        st.session_state.opinion.본문 = draft
        labels = risk_scan(draft)
        edits = suggest_edits(draft)
        log(f"리뷰 실행 — 리스크 {len(labels)}건, 제안 {len(edits)}건")

        st.markdown("#### 📛 규제 리스크 라벨")
        if not labels:
            st.success("표시할 리스크 키워드를 찾지 못했습니다.")
        for label, sev, why in labels:
            c = st.columns([1,1,8])
            with c[0]:
                if sev == "심각":
                    st.markdown('<span style="background:#EF4444;color:white;padding:4px 8px;border-radius:8px;">심각</span>', unsafe_allow_html=True)
                elif sev == "주의":
                    st.markdown('<span style="background:#F59E0B;color:white;padding:4px 8px;border-radius:8px;">주의</span>', unsafe_allow_html=True)
                else:
                    st.markdown('<span style="background:#64748B;color:white;padding:4px 8px;border-radius:8px;">정보</span>', unsafe_allow_html=True)
            with c[1]:
                st.write(f"**{label}**")
            with c[2]:
                st.write(why)

        st.markdown("#### ✍️ 수정 제안")
        for i, e in enumerate(edits, 1):
            st.markdown(f"- {i}. {e}")

        st.markdown("#### 🔎 인라인 하이라이트")
        st.markdown(highlight_risks_html(draft), unsafe_allow_html=True)

        # 라벨 내보내기 (CSV/XLSX)
        st.markdown("#### ⬇️ 리스크 라벨 내보내기")
        if labels:
            rows = [{"label": l, "severity": s, "rationale": w} for (l, s, w) in labels]
            csv = "label,severity,rationale\n" + "\n".join([f"{r['label']},{r['severity']},{r['rationale']}" for r in rows])
            st.download_button("CSV 다운로드", data=csv.encode("utf-8"), file_name="risk_labels.csv", mime="text/csv")
            if pd is not None:
                df = pd.DataFrame(rows)
                bio = io.BytesIO()
                with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
                    df.to_excel(writer, index=False, sheet_name="labels")
                st.download_button("Excel(.xlsx) 다운로드", data=bio.getvalue(), file_name="risk_labels.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.caption("내보낼 라벨이 없습니다.")

# =====================[ ④ 의견서 생성 ]=====================
elif flow == "④ 의견서 생성":
    st.subheader("④ 의견서 자동 생성 (멀티 에이전트 합성 데모)")

    title = st.text_input("의견서 제목", value=st.session_state.opinion.제목 or f"[{st.session_state.eng.법률분야}] {st.session_state.eng.목적} 의견서".strip())
    purpose = st.session_state.eng.목적 or st.text_input("자문 목적(없다면 입력)", "")
    domain = st.session_state.eng.법률분야 or st.selectbox("법률 분야 선택", ["개인정보","전자금융","노동","지식재산","공정거래","계약","기타"])

    if st.button("초안 생성"):
        outline = multi_agent_outline(domain, purpose, st.session_state.extras)
        body_parts = [f"## {sec}\n{txt}" for sec, txt in outline.items()]
        body = f"# {title}\n\n" + "\n\n".join(body_parts)
        labels = risk_scan(body)

        st.session_state.opinion = OpinionDoc(제목=title, 본문=body, 리스크라벨=labels)
        st.success("의견서 초안 생성 완료")
        log(f"의견서 초안 생성 — 라벨 {len(labels)}건")

        st.markdown(st.session_state.opinion.본문)

        st.markdown("#### 📛 자동 리스크 라벨")
        for label, sev, why in labels or []:
            c = st.columns([1,1,8])
            with c[0]:
                if sev == "심각":
                    st.markdown('<span style="background:#EF4444;color:white;padding:4px 8px;border-radius:8px;">심각</span>', unsafe_allow_html=True)
                elif sev == "주의":
                    st.markdown('<span style="background:#F59E0B;color:white;padding:4px 8px;border-radius:8px;">주의</span>', unsafe_allow_html=True)
                else:
                    st.markdown('<span style="background:#64748B;color:white;padding:4px 8px;border-radius:8px;">정보</span>', unsafe_allow_html=True)
            with c[1]:
                st.write(f"**{label}**")
            with c[2]:
                st.write(why)

# =====================[ ⑤ 출력/발송 ]=====================
elif flow == "⑤ 출력/발송":
    st.subheader("⑤ 의견서 출력 및 발송 준비")

    if not st.session_state.opinion.본문:
        st.warning("먼저 ③ 초안 리뷰 또는 ④ 의견서 생성에서 본문을 준비해 주세요.")
    else:
        st.markdown("#### 미리보기")
        st.markdown(st.session_state.opinion.본문)

        # Markdown 내보내기
        st.markdown("#### 📤 내보내기")
        md_bytes = st.session_state.opinion.본문.encode("utf-8")
        st.download_button("⬇️ Markdown(.md) 다운로드", data=md_bytes, file_name="opinion.md", mime="text/markdown")

        # DOCX 내보내기 (설치 시)
        if Document is not None:
            if st.button("DOCX 파일 만들기"):
                doc = Document()
                for line in st.session_state.opinion.본문.splitlines():
                    if line.startswith("# "):
                        doc.add_heading(line.replace("# ", ""), level=1)
                    elif line.startswith("## "):
                        doc.add_heading(line.replace("## ", ""), level=2)
                    else:
                        doc.add_paragraph(line)
                bio = io.BytesIO()
                doc.save(bio)
                st.download_button("⬇️ Word(.docx) 다운로드", data=bio.getvalue(), file_name="opinion.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.info("python-docx 미설치로 DOCX 버튼 숨김 (설치: pip install python-docx)")

        # 이메일 .eml 미리보기
        st.markdown("#### 📧 이메일 발송(미리보기 .eml 생성)")
        to_addr = st.text_input("받는사람 이메일", value="")
        from_addr = st.text_input("보내는사람 이메일", value="noreply@example.com")
        subject = st.text_input("제목", value=st.session_state.opinion.제목 or "법률 의견서")
        body_text = st.text_area("이메일 본문", value="의견서를 첨부드립니다.\n\n감사합니다.", height=120)

        if st.button("EML 미리보기 생성"):
            from email.message import EmailMessage
            msg = EmailMessage()
            msg["To"] = to_addr
            msg["From"] = from_addr
            msg["Subject"] = subject
            msg.set_content(body_text + "\n\n---\n" + st.session_state.opinion.본문)

            eml_bytes = msg.as_bytes()
            st.download_button("⬇️ EML 다운로드", data=eml_bytes, file_name="opinion.eml", mime="message/rfc822")
            st.success("이메일 파일(.eml)이 생성되었습니다. 로컬 메일 클라이언트에서 열어 전송하세요.")

        st.markdown("---")
        st.caption("브라우저 자동입력/오피스 자동화는 보안상 별도 워커(Selenium 등)에서 실행 권장. 본 데모는 파일 생성까지만 제공합니다.")
